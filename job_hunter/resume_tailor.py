"""
job_hunter/resume_tailor.py

Tailors an existing resume to a specific job listing: an OpenAI model
compares the candidate's original resume text against the pasted job
listing text, then reprioritizes/rewords the resume (summary, skills
order, experience bullets) to better match that role - without inventing
experience or skills that aren't already in the original resume.

The result is a structured JSON document (not raw prose) so it can be
rendered consistently to both .docx and .pdf regardless of how the
original resume file was formatted.
"""

import json
import os

from openai import OpenAI

# Fields the AI response must include for rendering to succeed.
REQUIRED_FIELDS = ["contact", "summary", "skills", "experience"]

_SYSTEM_PROMPT = """\
You are a resume-tailoring assistant. You are given (1) the full text of a \
candidate's existing resume and (2) the full text of a specific job listing \
they want to apply to. Rewrite/reprioritize the resume to better match this \
job listing, following these rules strictly:

- Do NOT invent skills, experience, employers, titles, dates, or \
  accomplishments that are not already present (even implicitly) in the \
  original resume. Only rephrase, reprioritize, and re-emphasize what's \
  really there.
- Reorder the skills list so the most relevant to THIS listing come first.
- Reword the professional summary to speak directly to this role/company.
- For each experience entry, rewrite the bullets to foreground the parts of \
  that job most relevant to this listing, using language/terminology that \
  mirrors the listing where honestly applicable. Keep the same jobs, \
  titles, companies, and dates as the original resume.
- If the listing asks for something the candidate's resume doesn't clearly \
  support, do not fake it - just don't over-emphasize that gap, and note it \
  in "candidate_notes" instead so the candidate can decide how to address it \
  (e.g. in a cover letter).

Return ONLY a single valid JSON object (no markdown, no commentary) matching \
exactly this schema:

{
  "contact": {
    "name": string,
    "location": string,
    "email": string,          // "" if not present in the original resume
    "phone": string,          // "" if not present in the original resume
    "linkedin": string        // "" if not present in the original resume
  },
  "target_role": string,       // job title from the listing, for the filename/header
  "target_company": string,    // company name from the listing, for the filename/header
  "summary": string,           // 2-4 sentence tailored professional summary
  "skills": [string],          // reordered/pruned from the original resume's skills
  "experience": [
    {
      "title": string,
      "company": string,
      "location": string,
      "dates": string,
      "bullets": [string]     // reworded/reprioritized bullets for this role
    }
  ],
  "education": [
    {"degree": string, "school": string, "dates": string}
  ],
  "certifications": [string],
  "candidate_notes": [string]  // gaps/mismatches worth addressing in a cover letter, or []
}
"""


def _validate_tailored(doc: dict):
    missing = [f for f in REQUIRED_FIELDS if f not in doc]
    if missing:
        raise ValueError(
            f"AI response is missing required field(s): {missing}. "
            "Try again, or check the model output for malformed JSON."
        )
    if not isinstance(doc.get("skills"), list) or not doc["skills"]:
        raise ValueError("AI response did not include any skills - can't render a tailored resume.")


def tailor_resume(resume_text: str, job_listing_text: str, model: str = None) -> dict:
    """
    Send the original resume text plus a pasted job listing to an OpenAI
    model and get back a structured, tailored resume document (see
    _SYSTEM_PROMPT for the schema). Raises if the API key is missing or the
    response doesn't match the required schema.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is not set. Add it to your .env file - "
            "get a key at https://platform.openai.com/api-keys"
        )

    model = model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    client = OpenAI(api_key=api_key)

    response = client.chat.completions.create(
        model=model,
        response_format={"type": "json_object"},
        temperature=0.3,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"Original resume text:\n\n{resume_text}\n\n"
                    f"---\n\nJob listing text:\n\n{job_listing_text}"
                ),
            },
        ],
    )

    raw = response.choices[0].message.content
    try:
        doc = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(f"AI response was not valid JSON: {e}\n\nRaw response:\n{raw}")

    _validate_tailored(doc)
    return doc


def render_docx(doc: dict, out_path: str):
    """Render a tailored resume dict to a clean, ATS-friendly .docx file."""
    from docx import Document
    from docx.shared import Pt

    d = Document()
    contact = doc.get("contact", {})

    title = d.add_heading(contact.get("name", "") or "Resume", level=1)
    title.alignment = 1

    contact_line = " | ".join(
        v for v in [contact.get("location"), contact.get("email"), contact.get("phone"), contact.get("linkedin")] if v
    )
    if contact_line:
        p = d.add_paragraph(contact_line)
        p.alignment = 1

    if doc.get("summary"):
        d.add_heading("Summary", level=2)
        d.add_paragraph(doc["summary"])

    if doc.get("skills"):
        d.add_heading("Skills", level=2)
        d.add_paragraph(" \u00b7 ".join(doc["skills"]))

    if doc.get("experience"):
        d.add_heading("Experience", level=2)
        for job in doc["experience"]:
            header = f"{job.get('title', '')} \u2014 {job.get('company', '')}"
            sub = " | ".join(v for v in [job.get("location"), job.get("dates")] if v)
            p = d.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            if sub:
                p.add_run(f"  ({sub})")
            for bullet in job.get("bullets", []):
                d.add_paragraph(bullet, style="List Bullet")

    if doc.get("education"):
        d.add_heading("Education", level=2)
        for edu in doc["education"]:
            line = f"{edu.get('degree', '')} \u2014 {edu.get('school', '')}"
            if edu.get("dates"):
                line += f" ({edu['dates']})"
            d.add_paragraph(line)

    if doc.get("certifications"):
        d.add_heading("Certifications", level=2)
        d.add_paragraph(" \u00b7 ".join(doc["certifications"]))

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    d.save(out_path)


def render_pdf(doc: dict, out_path: str):
    """Render a tailored resume dict to a clean .pdf file (no MS Word needed)."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

    contact = doc.get("contact", {})
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterSmall", parent=styles["Normal"], alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="NameHeading", parent=styles["Title"], alignment=TA_CENTER))

    story = [Paragraph(contact.get("name", "") or "Resume", styles["NameHeading"])]

    contact_line = " | ".join(
        v for v in [contact.get("location"), contact.get("email"), contact.get("phone"), contact.get("linkedin")] if v
    )
    if contact_line:
        story.append(Paragraph(contact_line, styles["CenterSmall"]))
    story.append(Spacer(1, 12))

    if doc.get("summary"):
        story.append(Paragraph("Summary", styles["Heading2"]))
        story.append(Paragraph(doc["summary"], styles["Normal"]))
        story.append(Spacer(1, 8))

    if doc.get("skills"):
        story.append(Paragraph("Skills", styles["Heading2"]))
        story.append(Paragraph(" &middot; ".join(doc["skills"]), styles["Normal"]))
        story.append(Spacer(1, 8))

    if doc.get("experience"):
        story.append(Paragraph("Experience", styles["Heading2"]))
        for job in doc["experience"]:
            header = f"<b>{job.get('title', '')} &mdash; {job.get('company', '')}</b>"
            sub = " | ".join(v for v in [job.get("location"), job.get("dates")] if v)
            if sub:
                header += f"  ({sub})"
            story.append(Paragraph(header, styles["Normal"]))
            bullets = job.get("bullets", [])
            if bullets:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(b, styles["Normal"])) for b in bullets],
                        bulletType="bullet",
                    )
                )
            story.append(Spacer(1, 6))

    if doc.get("education"):
        story.append(Paragraph("Education", styles["Heading2"]))
        for edu in doc["education"]:
            line = f"{edu.get('degree', '')} &mdash; {edu.get('school', '')}"
            if edu.get("dates"):
                line += f" ({edu['dates']})"
            story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 8))

    if doc.get("certifications"):
        story.append(Paragraph("Certifications", styles["Heading2"]))
        story.append(Paragraph(" &middot; ".join(doc["certifications"]), styles["Normal"]))

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    pdf = SimpleDocTemplate(out_path, pagesize=LETTER)
    pdf.build(story)
